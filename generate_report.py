"""
generate_report.py  –  CSUG CSC510 Module 6
--------------------------------------------
Generates the project documentation as a formatted .docx file.
Run once:  python generate_report.py
Output  :  CSUG_CSC510_M6_Naive_Bayes_Report.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── helpers ──────────────────────────────────────────────────────────────────

def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "4F81BD")
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def heading(doc, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return h


def body(doc, text: str, bold_prefix: str = ""):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
    p.add_run(text)
    return p


def bullet(doc, text: str, level: int = 0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    return p


def code_block(doc, text: str):
    for line in text.splitlines():
        p = doc.add_paragraph(line if line else " ")
        p.style = doc.styles["No Spacing"]
        p.paragraph_format.left_indent  = Inches(0.4)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        for run in p.runs:
            run.font.name = "Courier New"
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)
        pPr  = p._p.get_or_add_pPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "F2F2F2")
        pPr.append(shd)


def reference(doc, number: int, parts: tuple):
    p = doc.add_paragraph()
    p.paragraph_format.space_after       = Pt(3)
    p.paragraph_format.first_line_indent = Inches(-0.35)
    p.paragraph_format.left_indent       = Inches(0.35)
    r = p.add_run(f"[{number}]  ")
    r.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    if len(parts) == 3:
        author, title_text, rest = parts
        p.add_run(author)
        italic_run = p.add_run(title_text)
        italic_run.italic = True
        p.add_run(rest)
    else:
        p.add_run(parts[0])


# ── document ─────────────────────────────────────────────────────────────────

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

# ── TITLE PAGE ───────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("Naive Bayes Classifier\nUsing scikit-learn")
tr.font.size  = Pt(26)
tr.font.bold  = True
tr.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

doc.add_paragraph()

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run(
    "CSC510 – Module 6 Assignment\n"
    "Naive Bayes Classification with Laplacian Correction"
)
sr.font.size  = Pt(13)
sr.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

doc.add_paragraph()

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_p.add_run(
    f"Date: {datetime.date.today().strftime('%B %d, %Y')}"
).font.size = Pt(11)

add_horizontal_rule(doc)
doc.add_paragraph()

# ── 1. INTRODUCTION ──────────────────────────────────────────────────────────
heading(doc, "1. Introduction", 1)
body(doc,
     "Machine learning classification is a fundamental task in data science and artificial "
     "intelligence. Among the many classification algorithms available, the Naive Bayes "
     "classifier stands out for its simplicity, speed, and surprisingly strong performance "
     "across a wide variety of domains — including text classification, spam filtering, "
     "medical diagnosis, and sentiment analysis [1].")
body(doc,
     "This project implements a fully functional Naive Bayes classifier in Python using "
     "the scikit-learn machine learning library [2]. The classifier is built upon Bayes' "
     "Theorem and assumes conditional independence among features — the assumption that "
     "gives the algorithm its 'naive' label. Despite this simplification, Naive Bayes "
     "classifiers often rival more complex models in performance, particularly when "
     "training data is limited [3].")
body(doc,
     "The classic Play Tennis dataset is used as the training corpus. This 14-sample "
     "dataset, introduced by Quinlan [4], is a widely used pedagogical example that "
     "contains four categorical weather features (Outlook, Temperature, Humidity, Wind) "
     "and a binary label indicating whether tennis was played on a given day.")

# ── 2. BACKGROUND ────────────────────────────────────────────────────────────
heading(doc, "2. Theoretical Background", 1)

heading(doc, "2.1  Bayes' Theorem", 2)
body(doc,
     "Bayes' Theorem describes the probability of an event based on prior knowledge of "
     "conditions related to that event. Formally [1][5]:")
code_block(doc, "P(C | X) = P(X | C) × P(C) / P(X)")
body(doc,
     "Where:\n"
     "  • P(C | X)  — posterior probability of class C given features X\n"
     "  • P(X | C)  — likelihood of observing features X given class C\n"
     "  • P(C)      — prior probability of class C\n"
     "  • P(X)      — marginal probability of observing X (normalisation constant)")

heading(doc, "2.2  The Naive Independence Assumption", 2)
body(doc,
     "The 'naive' in Naive Bayes refers to the assumption that all features are "
     "conditionally independent given the class label [1]. This simplifies the "
     "likelihood calculation to a product of individual feature probabilities:")
code_block(doc, "P(X | C) = P(x1 | C) × P(x2 | C) × ... × P(xn | C)")
body(doc,
     "Although this assumption rarely holds in practice, Naive Bayes classifiers have "
     "been shown to perform competitively against more sophisticated algorithms, "
     "especially on high-dimensional data [6].")

heading(doc, "2.3  Categorical Naive Bayes", 2)
body(doc,
     "This project employs scikit-learn's CategoricalNB model, which is specifically "
     "designed for discrete, categorical features. Unlike GaussianNB (which assumes a "
     "Gaussian distribution of continuous features) or BernoulliNB (which assumes binary "
     "features), CategoricalNB models the frequency distribution of each category per "
     "class using a categorical distribution [2][7].")

heading(doc, "2.4  Laplacian (Additive) Smoothing", 2)
body(doc,
     "A fundamental problem in Naive Bayes classification is the zero-probability "
     "problem: if a feature value never co-occurs with a class in the training data, "
     "its estimated probability is zero, which zeroes out the entire posterior "
     "regardless of all other evidence [8].")
body(doc,
     "Laplacian smoothing (also called additive smoothing) resolves this by adding a "
     "small pseudo-count α (typically 1) to every feature-value count:")
code_block(doc, "P(value | class) = (count(value, class) + α) / (count(class) + α × k)")
body(doc,
     "Where k is the number of unique values for that feature. With α = 1, no "
     "probability is ever exactly zero [8][9]. The scikit-learn CategoricalNB model "
     "exposes this directly via the alpha parameter.")

# ── 3. PROJECT OVERVIEW ──────────────────────────────────────────────────────
heading(doc, "3. Project Overview", 1)
body(doc,
     "The project is structured as a Python package following the Single Responsibility "
     "Principle. Each module has a distinct, clearly defined role, promoting "
     "maintainability, testability, and extensibility.")

heading(doc, "3.1  Project Structure", 2)
code_block(doc,
"""CSUG_CSC510_M6/
├── main.py                    Entry point
├── requirements.txt           Dependencies
├── generate_report.py         This report generator
└── naive_bayes/               Package
    ├── __init__.py            Public API
    ├── logger.py              Centralised logging
    ├── dataset.py             Dataset definition
    ├── tables.py              Frequency & Likelihood tables
    ├── model.py               scikit-learn model wrapper
    ├── walkthrough.py         Manual Bayes math
    └── classifier.py         Pipeline orchestrator""")

heading(doc, "3.2  Module Descriptions", 2)

tbl = doc.add_table(rows=1, cols=2)
tbl.style = "Table Grid"
hdr = tbl.rows[0].cells
hdr[0].text = "Module"
hdr[1].text = "Responsibility"
for cell in hdr:
    set_cell_bg(cell, "1F497D")
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

rows_data = [
    ("logger.py",      "Provides get_logger() — a single configuration point for all "
                       "log formatting, levels, and handlers across the package."),
    ("dataset.py",     "Defines PlayTennisDataset, which owns the raw DataFrame, "
                       "feature list, target column, and class labels."),
    ("tables.py",      "FrequencyTable builds the raw count table and the "
                       "Laplacian-smoothed likelihood table, flagging zero-count cells."),
    ("model.py",       "NaiveBayesModel wraps OrdinalEncoder + CategoricalNB. Handles "
                       "encoding, training, and single-sample prediction."),
    ("walkthrough.py", "PosteriorWalkthrough manually computes and prints the full "
                       "step-by-step Bayes' Theorem calculation for a given sample."),
    ("classifier.py",  "NaiveBayesClassifier is the top-level orchestrator that "
                       "wires all components together and runs the pipeline."),
    ("main.py",        "Clean entry point. Defines test_samples and custom_test, "
                       "instantiates NaiveBayesClassifier, and calls run()."),
]
for mod, desc in rows_data:
    row = tbl.add_row().cells
    row[0].text = mod
    row[1].text = desc
    for para in row[0].paragraphs:
        for run in para.runs:
            run.font.bold = True
            run.font.name = "Courier New"
            run.font.size = Pt(9)

doc.add_paragraph()

# ── 4. DATASET ───────────────────────────────────────────────────────────────
heading(doc, "4. Dataset", 1)
body(doc,
     "The Play Tennis dataset, originally used by Quinlan [4] to illustrate decision "
     "tree induction and later widely adopted in Naive Bayes tutorials [3], consists of "
     "14 training samples. Each sample records weather conditions on a given day and "
     "whether tennis was played.")
body(doc, "Features and their possible values:")
bullet(doc, "Outlook      : Sunny, Overcast, Rainy")
bullet(doc, "Temperature  : Hot, Mild, Cool")
bullet(doc, "Humidity     : High, Normal")
bullet(doc, "Wind         : Weak, Strong")
bullet(doc, "Play (label) : Yes (9 samples), No (5 samples)")
body(doc,
     "The dataset contains a zero-count cell: Overcast never appears with class 'No'. "
     "This makes it an ideal example for demonstrating Laplacian correction, as without "
     "smoothing P(Overcast | No) = 0 would incorrectly eliminate the 'No' class.")

# ── 5. CLASSIFICATION PIPELINE ───────────────────────────────────────────────
heading(doc, "5. Classification Pipeline", 1)

heading(doc, "5.1  Step 1 – Frequency Table", 2)
body(doc,
     "The dataset is converted into a frequency table by counting how many times each "
     "feature value appears for each class label. For example, for the Outlook feature:")
code_block(doc,
"""Outlook   | No | Yes | Total
----------+----+-----+------
Overcast  |  0 |   4 |     4   ← zero count (Laplacian fix needed)
Rainy     |  2 |   3 |     5
Sunny     |  3 |   2 |     5""")

heading(doc, "5.2  Step 2 – Likelihood Table", 2)
body(doc,
     "Each raw count is converted to a smoothed conditional probability using the "
     "Laplacian correction formula. The *** marker highlights zero-count cells:")
code_block(doc,
"""P(Overcast | No)  = (0 + 1) / (5 + 1×3) = 0.1250  ***
P(Overcast | Yes) = (4 + 1) / (9 + 1×3) = 0.4167
P(Rainy    | No)  = (2 + 1) / (5 + 1×3) = 0.3750
P(Sunny    | No)  = (3 + 1) / (5 + 1×3) = 0.5000""")

heading(doc, "5.3  Step 3 – Prior Probabilities", 2)
code_block(doc,
"""P(No)  = 5 / 14 = 0.3571
P(Yes) = 9 / 14 = 0.6429""")

heading(doc, "5.4  Step 4 – Model Training", 2)
body(doc,
     "Features are ordinally encoded using scikit-learn's OrdinalEncoder and the "
     "CategoricalNB model is trained with alpha=1.0 (Laplacian correction). "
     "The model internally stores log-probabilities for numerically stable predictions [2].")

heading(doc, "5.5  Step 5 – Posterior Probability Prediction", 2)
code_block(doc,
"""log P(No  | Sunny, Hot, High, Weak)
  = log(0.3571) + log(0.5000) + log(0.3750) + log(0.7143) + log(0.4286)
  = -3.8874

log P(Yes | Sunny, Hot, High, Weak) = -4.6780

→ Normalised:  P(No | features)  = 0.6880
               P(Yes | features) = 0.3120
→ Prediction:  No  (do not play tennis)""")

heading(doc, "5.6  Step 6 – Manual Walkthrough", 2)
body(doc,
     "The PosteriorWalkthrough class independently recomputes the posterior from "
     "scratch using raw frequency counts, providing transparent step-by-step "
     "verification of the model's output without relying on scikit-learn internals.")

# ── 6. IMPLEMENTATION DETAILS ────────────────────────────────────────────────
heading(doc, "6. Key Implementation Details", 1)

heading(doc, "6.1  Logging Strategy", 2)
body(doc,
     "All modules use Python's built-in logging framework via a centralised "
     "get_logger(__name__) factory in logger.py [10]:")
bullet(doc, "INFO  — major pipeline events (training complete, pipeline start/end)")
bullet(doc, "DEBUG — per-cell probability calculations, sample predictions")
bullet(doc, "ERROR — guard clauses (e.g., predict() called before train())")

heading(doc, "6.2  Modular Design", 2)
body(doc,
     "The package follows the Single Responsibility Principle from SOLID design [11]: "
     "each class does exactly one job. Swapping the dataset requires only a change to "
     "dataset.py; changing the smoothing strategy requires only a change to _smooth() "
     "in tables.py and the alpha parameter in model.py.")

heading(doc, "6.3  Zero-Probability Protection", 2)
body(doc,
     "Laplacian correction is applied in two independent places: (1) manually in "
     "FrequencyTable._smooth() for display purposes, and (2) automatically by "
     "CategoricalNB(alpha=1.0) during training. This dual implementation ensures "
     "the displayed likelihood table exactly matches the model's internal probabilities, "
     "providing a verifiable audit trail [8].")

# ── 7. TESTING ───────────────────────────────────────────────────────────────
heading(doc, "7. Testing Instructions", 1)

heading(doc, "7.1  Setup", 2)
code_block(doc,
"""# Install dependencies
pip install scikit-learn pandas tabulate numpy python-docx

# Run the classifier
python main.py

# Generate this report
python generate_report.py""")

heading(doc, "7.2  Running Custom Predictions", 2)
body(doc,
     "Open main.py and edit the custom_test dictionary inside main():")
code_block(doc,
"""custom_test = {
    "Outlook":     "Rainy",     # "Sunny", "Overcast", "Rainy"
    "Temperature": "Cool",      # "Hot", "Mild", "Cool"
    "Humidity":    "High",      # "High", "Normal"
    "Wind":        "Strong",    # "Weak", "Strong"
}""")
body(doc, "Then re-run:  python main.py")

heading(doc, "7.3  Interpreting the Output", 2)
tbl2 = doc.add_table(rows=1, cols=2)
tbl2.style = "Table Grid"
hdr2 = tbl2.rows[0].cells
hdr2[0].text = "Output Section"
hdr2[1].text = "What It Shows"
for cell in hdr2:
    set_cell_bg(cell, "1F497D")
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

output_rows = [
    ("Dataset",               "Full 14-row training table"),
    ("Frequency Table",       "Raw counts per feature-value per class"),
    ("Likelihood Table",      "Laplacian-smoothed probabilities; *** flags zero counts"),
    ("Prior Probabilities",   "P(No) and P(Yes) from training data"),
    ("Training Summary",      "Model type and alpha confirmation"),
    ("Posterior Predictions", "P(No), P(Yes), and prediction for 5 test samples"),
    ("Manual Walkthrough",    "Step-by-step Bayes math for Sample 1"),
    ("Custom Test",           "Prediction for the user-defined custom_test input"),
]
for sec, desc in output_rows:
    row = tbl2.add_row().cells
    row[0].text = sec
    row[1].text = desc
    for para in row[0].paragraphs:
        for run in para.runs:
            run.font.bold = True

doc.add_paragraph()

# ── 8. RESULTS ───────────────────────────────────────────────────────────────
heading(doc, "8. Sample Results", 1)
body(doc,
     "The table below shows the posterior probabilities and predictions output by the "
     "trained CategoricalNB model for five test samples:")

tbl3 = doc.add_table(rows=1, cols=7)
tbl3.style = "Table Grid"
headers3 = ["Outlook", "Temperature", "Humidity", "Wind", "P(No)", "P(Yes)", "Prediction"]
for i, h in enumerate(headers3):
    cell = tbl3.rows[0].cells[i]
    cell.text = h
    set_cell_bg(cell, "1F497D")
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

results = [
    ("Sunny",    "Hot",  "High",   "Weak",   "0.6880", "0.3120", "No"),
    ("Overcast", "Cool", "Normal", "Strong", "0.0810", "0.9190", "Yes"),
    ("Rainy",    "Mild", "High",   "Strong", "0.6345", "0.3655", "No"),
    ("Sunny",    "Cool", "Normal", "Weak",   "0.2013", "0.7987", "Yes"),
    ("Rainy",    "Hot",  "Normal", "Weak",   "0.2209", "0.7791", "Yes"),
]
for i, r in enumerate(results):
    row = tbl3.add_row().cells
    for j, val in enumerate(r):
        row[j].text = val
        if i % 2 == 0:
            set_cell_bg(row[j], "DCE6F1")
        if j == 6:
            for para in row[j].paragraphs:
                for run in para.runs:
                    run.font.bold = True
                    run.font.color.rgb = (
                        RGBColor(0x00, 0x70, 0x00) if val == "Yes"
                        else RGBColor(0xC0, 0x00, 0x00)
                    )

doc.add_paragraph()

# ── 9. DISCUSSION ────────────────────────────────────────────────────────────
heading(doc, "9. Discussion", 1)
body(doc, "This project successfully demonstrates all four core requirements:")
bullet(doc, "Frequency Table     — raw groupby counts per class")
bullet(doc, "Likelihood Table    — Laplacian-smoothed probabilities (α=1)")
bullet(doc, "Posterior Probability — calculated per class and normalised")
bullet(doc, "Zero Probability Fix  — Laplacian correction in both display tables "
            "and the scikit-learn model (alpha=1.0)")
body(doc,
     "The modular design separates concerns clearly across six dedicated classes. "
     "The logger provides a full audit trail of every calculation. The manual walkthrough "
     "verifies the model's scikit-learn predictions from first principles.")

# ── 10. CONCLUSION ───────────────────────────────────────────────────────────
heading(doc, "10. Conclusion", 1)
body(doc,
     "This project demonstrates a complete, production-quality Naive Bayes classifier "
     "implemented in Python using scikit-learn. The classifier correctly computes "
     "frequency tables, likelihood tables with Laplacian correction, and posterior "
     "probabilities for both test samples and user-defined custom inputs. The codebase "
     "is organised as a modular Python package with centralised logging, clean "
     "separation of concerns, and comprehensive inline documentation.")
body(doc,
     "Naive Bayes classifiers remain a relevant and widely used technique in modern "
     "machine learning pipelines [6]. Their efficiency and interpretability make them "
     "a valuable tool for both production applications and educational purposes.")

# ── REFERENCES ───────────────────────────────────────────────────────────────
add_horizontal_rule(doc)
heading(doc, "References", 1)

refs = [
    ("Mitchell, T. M. (1997). ",
     "Machine Learning. ",
     "McGraw-Hill. Chapter 6: Bayesian Learning. ISBN 0-07-042807-7."),

    ("Pedregosa, F., Varoquaux, G., Gramfort, A., et al. (2011). ",
     "Scikit-learn: Machine learning in Python. ",
     "Journal of Machine Learning Research, 12, 2825–2830. "
     "https://jmlr.org/papers/v12/pedregosa11a.html"),

    ("Rish, I. (2001). ",
     "An empirical study of the naive Bayes classifier. ",
     "IJCAI 2001 Workshop on Empirical Methods in AI, 3(22), 41–46."),

    ("Quinlan, J. R. (1986). ",
     "Induction of decision trees. ",
     "Machine Learning, 1(1), 81–106. https://doi.org/10.1007/BF00116251"),

    ("Bayes, T., & Price, R. (1763). ",
     "An essay towards solving a problem in the doctrine of chances. ",
     "Philosophical Transactions of the Royal Society of London, 53, 370–418."),

    ("Domingos, P., & Pazzani, M. (1997). ",
     "On the optimality of the simple Bayesian classifier under zero-one loss. ",
     "Machine Learning, 29(2-3), 103–130. https://doi.org/10.1023/A:1007413511361"),

    ("Zhang, H. (2004). ",
     "The optimality of naive Bayes. ",
     "Proceedings of FLAIRS 2004, 3(1), 562–567. AAAI Press."),

    ("Lidstone, G. J. (1920). ",
     "Note on the general case of the Bayes-Laplace formula. ",
     "Transactions of the Faculty of Actuaries, 8, 182–192."),

    ("Manning, C. D., Raghavan, P., & Schutze, H. (2008). ",
     "Introduction to Information Retrieval. ",
     "Cambridge University Press. https://nlp.stanford.edu/IR-book/"),

    ("Python Software Foundation. (2024). ",
     "logging — Logging facility for Python. ",
     "Python 3 Documentation. https://docs.python.org/3/library/logging.html"),

    ("Martin, R. C. (2003). ",
     "Agile Software Development: Principles, Patterns, and Practices. ",
     "Prentice Hall. Chapter 8: The Single Responsibility Principle. "
     "ISBN 0-13-597444-5."),
]

for i, parts in enumerate(refs, start=1):
    reference(doc, i, parts)

# ── SAVE ─────────────────────────────────────────────────────────────────────
output_path = r"C:\Users\prose\PycharmProjects\CSUG_CSC510_M5\CSUG_CSC510_M6_Naive_Bayes_Report.docx"
doc.save(output_path)
print(f"Report saved to: {output_path}")

